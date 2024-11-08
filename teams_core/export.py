import os
from docx import Document
import pypandoc
import json
import datetime 
from docx.shared import Inches

def add_test_case_to_doc(doc, testcase, test_count=1):
    # Procedural Steps
    doc.add_heading(f"3.{test_count} Test Case: " + testcase.name, level=3)

    # Parse the JSON content
    content = None
    if testcase.content is not None and len(testcase.content.strip()) >0:
        try:
            content = json.loads(testcase.content)
        except:
            print('Error in parsing content')
            print(testcase.content)
            content = {'proceduralSteps': {},
                       'description': ''}
    else:
        content = {'proceduralSteps': {},
                    'description': ''}

    # Create an overall table with initial rows and columns
    overall_table = doc.add_table(rows=3, cols=2)
    overall_table.autofit = False
    overall_table.columns[0].width = Inches(1.5)
    overall_table.columns[1].width = Inches(4.5)

    # Add data for Test Case Number and Name
    current_row = 0
    overall_table.cell(0, 0).text = "Test Case Number:"
    overall_table.cell(0, 1).text = f"{testcase.oid}"
    overall_table.cell(1, 0).text = "Test Case Name:"
    overall_table.cell(1, 1).text = f"{testcase.name}"
    overall_table.cell(2, 0).text = "Test Case Description:"
    overall_table.cell(2, 1).text = f"{content['description']}"

    # Merge cells for Procedural Steps header
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Procedural Steps"

    # Add headers for Procedural Steps
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).text = "Step"
    overall_table.cell(current_row, 1).text = "Description"

    # Populate the procedural steps
    for step, details in content['proceduralSteps'].items():
        for detail in details:
            current_row += 1
            overall_table.add_row()
            overall_table.cell(current_row, 0).text = step
            overall_table.cell(current_row, 1).text = detail['stepDescription']

    # Merge cells for Environment/Setup header
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Environment/Setup"

    # Add placeholder text for Environment/Setup
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Specify environment/setup details here..."

    # Add Hardware, Software, and Other details
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).text = "Hardware"
    overall_table.cell(current_row, 1).text = "Specify hardware here..."
    
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).text = "Software"
    overall_table.cell(current_row, 1).text = "Specify software here..."
    
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).text = "Other"
    overall_table.cell(current_row, 1).text = "Specify other details here..."

    # Merge cells for Intercase Dependencies header
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Intercase Dependencies"


    # Add placeholder text for Intercase Dependencies
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Specify intercase dependencies here..."

    # Merge cells for Specifications header
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).merge(overall_table.cell(current_row, 1))
    overall_table.cell(current_row, 0).text = "Specifications"

    # Add headers for Specifications
    current_row += 1
    overall_table.add_row()
    overall_table.cell(current_row, 0).text = "Input"
    # overall_table.cell(current_row, 1).split(1, 3)
    overall_table.cell(current_row, 1).text = "Expected Output"
    # overall_table.cell(current_row, 2).text = "Actual Output"
    # Example specs (replace this with your own)

    # Populate the input and expected output based on procedural steps
    for step, details in content['proceduralSteps'].items():
        for detail in details:
            if len(detail['stepDescription'])==0 and len(detail['expectedOutput'])==0:
                continue
            row_cells = overall_table.add_row().cells
            row_cells[0].text = detail['stepDescription']  # Input description
            row_cells[1].text = detail['expectedOutput']  # Expected output
            # row_cells[2].text = ""  # Leave Actual Output blank


def generate_docx(testcase=None, testsuite=None):
    docx_filename = f"testcase_{testcase.id}.docx" if testsuite is None else f"testsuite_{testsuite.name.replace(' ', '_')}.docx"
    docx_path = f"/tmp/{docx_filename}"  # Store it temporarily
    # Create a DOCX document
    doc = Document()

    # Title
    doc.add_heading('Functional Test Case Document', level=1)
    if testsuite is not None:
        doc.add_heading(f"Document Code:", level=2)
        doc.add_heading(f"Test Case Name: {testsuite.name}", level=2)
        doc.add_heading(f"Version: ", level=2)
    else:
        doc.add_heading(f"Document Code: {testcase.oid}", level=2)
        doc.add_heading(f"Test Case Name: {testcase.name}", level=2)
        doc.add_heading(f"Version: {testcase.version}", level=2)

    # Approval Block
    doc.add_heading("Approval", level=3)
    doc.add_paragraph("Prepared by: [Your Name]")
    doc.add_paragraph("Approved by: [Approver Name]")
    doc.add_paragraph("Date: [Approval Date]")

    # Revision Chart
    doc.add_heading("Revision History", level=3)
    revision_table = doc.add_table(rows=1, cols=3)
    hdr_cells = revision_table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Description'

    # Add a row for the current version
    row_cells = revision_table.add_row().cells
    if testsuite is not None:
        row_cells[0].text = ''
    else:
        row_cells[0].text = testcase.version
    row_cells[1].text = datetime.datetime.now().strftime('%Y-%m-%d')
    row_cells[2].text = 'Initial creation'

    # Introduction Section
    doc.add_heading("1. Introduction", level=2)
    doc.add_heading("1.1 Purpose", level=3)
    if testsuite is not None:
        doc.add_paragraph(testsuite.content)  # Customize as needed
    else:
        doc.add_paragraph("The tests cover the following aspects...")  # Customize as needed
    doc.add_heading("1.2 Scope", level=3)
    doc.add_paragraph("The tests cover the following aspects...")  # Customize as needed
    doc.add_heading("1.3 Outline", level=3)
    doc.add_paragraph("This document includes the following sections...")  # Customize as needed
    doc.add_heading("1.4 Reference Material", level=3)
    doc.add_paragraph("List any references or related documents...")  # Customize as needed
    doc.add_heading("1.5 Definitions and Acronyms", level=3)
    doc.add_paragraph("Provide definitions and acronyms used in this document...")  # Customize as needed

    # Test Setup and Basic Configuration (Section 2)
    doc.add_heading("2. Test Setup and Basic Configuration", level=2)

    # Test Cases Section
    doc.add_heading("3. Test Cases and Reports", level=2)

    if testsuite is not None:
        ii = 1
        for testcase in testsuite.testcase_set.all():
            add_test_case_to_doc(doc, testcase, ii)
            ii += 1
    else:
        add_test_case_to_doc(doc, testcase)
    doc.save(docx_path)
    return [docx_filename, docx_path]



def generate_pdf(testcase, testsuite=None):
    # First, generate DOCX
    docx_filename, docx_path = generate_docx(testcase)

    # Convert DOCX to PDF
    pdf_filename = f"testcase_{testcase.id}.pdf"
    pdf_path = f"/tmp/{pdf_filename}"

    try:
        # Convert DOCX to PDF using pypandoc
        output = pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
    except OSError as e:
        return None, f"Error converting to PDF: {str(e)}"
    
    return pdf_filename, pdf_path
